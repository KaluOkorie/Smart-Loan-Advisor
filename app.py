"""
Smart Credit Risk Advisor
Professional Loan Eligibility Assessment System
"""

import streamlit as st
import pandas as pd
import joblib
import plotly.graph_objects as go
import plotly.express as px
from datetime import datetime
import base64
import numpy as np
import io
import time
import shap
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import plotly.io as pio

# ---------------------------------------------------------
# SETUP & CONFIGURATION
# ---------------------------------------------------------
st.set_page_config(
    page_title="Credit Risk Assessment System",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Professional CSS
st.markdown("""
<style>
    .main-header {
        font-size: 2.2rem;
        font-weight: 700;
        margin-bottom: 1.2rem;
        color: #1a237e;
    }
    
    .sub-header {
        font-size: 1.4rem;
        font-weight: 600;
        margin-top: 1.5rem;
        margin-bottom: 1rem;
        color: #283593;
    }
    
    .metric-container {
        background-color: #f5f7fa;
        padding: 1.2rem;
        border-radius: 8px;
        border-left: 4px solid #3f51b5;
        margin: 0.5rem 0;
    }
    
    .success-container {
        background-color: rgba(76, 175, 80, 0.1);
        padding: 1rem;
        border-radius: 8px;
        border: 1px solid #4caf50;
        margin: 0.5rem 0;
    }
    
    .warning-container {
        background-color: rgba(255, 152, 0, 0.1);
        padding: 1rem;
        border-radius: 8px;
        border: 1px solid #ff9800;
        margin: 0.5rem 0;
    }
    
    .info-container {
        background-color: rgba(33, 150, 243, 0.1);
        padding: 1rem;
        border-radius: 8px;
        border: 1px solid #2196f3;
        margin: 0.5rem 0;
    }
    
    .stMetric {
        min-height: 100px;
    }
    
    .stMetric > div {
        min-height: 85px;
    }
    
    .stMetric label {
        font-size: 0.9rem !important;
        font-weight: 600 !important;
        color: #455a64 !important;
    }
    
    .sidebar-section {
        background-color: #f8f9fa;
        padding: 1rem;
        border-radius: 8px;
        margin-bottom: 1.5rem;
        border: 1px solid #e0e0e0;
    }
    
    .disclaimer-box {
        background-color: #fff3e0;
        padding: 1rem;
        border-radius: 6px;
        border-left: 4px solid #ff9800;
        margin: 1rem 0;
        font-size: 0.85rem;
        line-height: 1.4;
    }
</style>
""", unsafe_allow_html=True)

# ---------------------------------------------------------
# INITIALIZE SESSION STATE
# ---------------------------------------------------------
if 'last_submission' not in st.session_state:
    st.session_state.last_submission = 0
if 'results' not in st.session_state:
    st.session_state.results = None
if 'charts' not in st.session_state:
    st.session_state.charts = {}

# ---------------------------------------------------------
# UK FINANCIAL STANDARDS
# ---------------------------------------------------------
CONVERSION_RATE = 100
UK_AVERAGE_SALARY = 35000
UK_AVERAGE_HOUSE_PRICE = 250000
UK_ASSET_MULTIPLIER = 5

# ---------------------------------------------------------
# FEATURE ENGINEERING FUNCTIONS
# ---------------------------------------------------------
def create_credit_features(df):
    """Create engineered features for credit assessment"""
    df_engineered = df.copy()
    
    # Monthly calculations
    monthly_income = df_engineered['income_annum'] / 12
    monthly_loan_payment = (df_engineered['loan_amount'] / df_engineered['loan_term']) / 12
    
    # Key ratios
    df_engineered['monthly_payment_ratio'] = (monthly_loan_payment / monthly_income) * 100
    df_engineered['asset_coverage'] = (df_engineered['total_assets'] / df_engineered['loan_amount']) * 100
    df_engineered['loan_to_income'] = (df_engineered['loan_amount'] / df_engineered['income_annum']) * 100
    
    # Credit score categories
    def categorize_credit_score(score):
        if score >= 900:
            return 'Excellent'
        elif score >= 800:
            return 'Very Good'
        elif score >= 700:
            return 'Good'
        elif score >= 600:
            return 'Fair'
        else:
            return 'Needs Improvement'
    
    df_engineered['credit_category'] = df_engineered['credit_score'].apply(categorize_credit_score)
    
    # Stability score
    df_engineered['stability_score'] = 0
    df_engineered.loc[df_engineered['self_employed'] == 'No', 'stability_score'] += 30
    df_engineered.loc[df_engineered['education'] == 'Graduate', 'stability_score'] += 20
    df_engineered.loc[df_engineered['no_of_dependents'] <= 2, 'stability_score'] += 10
    
    return df_engineered

def calculate_matching_score(row):
    """Calculate matching score (0-100) based on credit profile"""
    score = 0
    
    # 1. Credit Score Component
    if row['credit_score'] >= 900:
        score += 40
    elif row['credit_score'] >= 800:
        score += 35
    elif row['credit_score'] >= 700:
        score += 30
    elif row['credit_score'] >= 600:
        score += 20
    else:
        score += 10
    
    # 2. Monthly Payment Affordability
    if row['monthly_payment_ratio'] <= 25:
        score += 25
    elif row['monthly_payment_ratio'] <= 35:
        score += 20
    elif row['monthly_payment_ratio'] <= 45:
        score += 15
    elif row['monthly_payment_ratio'] <= 55:
        score += 10
    else:
        score += 5
    
    # 3. Asset Security
    if row['asset_coverage'] >= 250:
        score += 20
    elif row['asset_coverage'] >= 175:
        score += 16
    elif row['asset_coverage'] >= 125:
        score += 12
    elif row['asset_coverage'] >= 75:
        score += 8
    else:
        score += 4
    
    # 4. Stability Factors
    score += min(row['stability_score'], 15)
    
    # 5. Risk adjustments
    if row['self_employed'] == 'Yes' and row['income_annum'] < 30000:
        score -= 10
    
    if row['no_of_dependents'] > 3:
        score -= 5
    
    if row['loan_term'] > 7 and row['loan_amount'] < 150000:
        score -= 3
    
    return min(max(score, 0), 100)

def generate_detailed_recommendations(score, features, applicant_data, prediction):
    """Generate personalized, actionable credit risk recommendations"""
    recommendations = []
    
    # Credit Score Analysis
    credit_score = features.get('credit_score', 0)
    if credit_score >= 800:
        recommendations.append({
            "title": "Excellent Credit Standing",
            "message": f"Credit score of {credit_score} is in the top tier for lending.",
            "actions": [
                "Qualify for optimal interest rates",
                "Maintain credit utilization below 25%",
                "Continue current payment patterns"
            ],
            "box_type": "success"
        })
    elif credit_score >= 700:
        recommendations.append({
            "title": "Good Credit Profile",
            "message": f"Score of {credit_score} meets standard lending requirements.",
            "actions": [
                "Aim for 750+ to access premium rates",
                "Review credit report quarterly",
                "Limit new credit applications"
            ],
            "box_type": "success"
        })
    elif credit_score >= 600:
        recommendations.append({
            "title": "Credit Improvement Opportunity",
            "message": f"Score of {credit_score} requires attention for optimal terms.",
            "actions": [
                "Register on electoral roll if applicable",
                "Reduce credit card balances below 30%",
                "Establish consistent payment history"
            ],
            "box_type": "warning"
        })
    else:
        recommendations.append({
            "title": "Credit Building Required",
            "message": f"Score of {credit_score} indicates significant risk factors.",
            "actions": [
                "Obtain comprehensive credit reports",
                "Address any delinquencies immediately",
                "Build 12-month positive payment history"
            ],
            "box_type": "warning"
        })
    
    # Payment Affordability Analysis
    payment_ratio = features.get('monthly_payment_ratio', 0)
    monthly_payment = applicant_data['loan_amount'] / (applicant_data['loan_term'] * 12)
    
    if payment_ratio <= 30:
        recommendations.append({
            "title": "Strong Payment Capacity",
            "message": f"Monthly payment of £{monthly_payment:.0f} represents {payment_ratio:.1f}% of income.",
            "actions": [
                "Maintain current debt-to-income ratio",
                "Consider shorter terms for interest savings",
                "Monitor changes in income stability"
            ],
            "box_type": "success"
        })
    elif payment_ratio <= 40:
        recommendations.append({
            "title": "Acceptable Payment Load",
            "message": f"Payment of £{monthly_payment:.0f} is {payment_ratio:.1f}% of income.",
            "actions": [
                "Maintain emergency fund coverage",
                "Review discretionary expenses quarterly",
                "Consider longer terms for flexibility"
            ],
            "box_type": "info"
        })
    else:
        recommendations.append({
            "title": "Elevated Payment Burden",
            "message": f"At {payment_ratio:.1f}% of income, payment capacity is constrained.",
            "actions": [
                f"Reduce requested amount by £{applicant_data['loan_amount'] * 0.1:.0f}",
                "Extend loan term to improve affordability",
                "Increase income sources or reduce expenses"
            ],
            "box_type": "warning"
        })
    
    # Model Prediction Specific Recommendations
    if prediction == 1:
        recommendations.append({
            "title": "Model Assessment: Low Risk",
            "message": "Machine learning model identifies low default probability.",
            "actions": [
                "Proceed with formal application",
                "Document all income sources thoroughly",
                "Prepare 6 months of bank statements"
            ],
            "box_type": "success"
        })
    else:
        recommendations.append({
            "title": "Model Assessment: Elevated Risk",
            "message": "Model indicates elevated default risk based on profile characteristics.",
            "actions": [
                "Review and improve credit profile factors",
                "Consider reducing requested loan amount",
                "Provide additional collateral if available"
            ],
            "box_type": "warning"
        })
    
    # Overall Score-Based Recommendation
    if score >= 85:
        recommendations.append({
            "title": "Premium Credit Profile",
            "message": f"Matching score of {score}/100 indicates excellent creditworthiness.",
            "actions": [
                "Apply to multiple lenders for competitive terms",
                "Expect accelerated processing timeline",
                "Document all assets comprehensively"
            ],
            "box_type": "success"
        })
    elif score >= 70:
        recommendations.append({
            "title": "Strong Credit Position",
            "message": f"Score of {score}/100 suggests high approval probability.",
            "actions": [
                "Complete application with full documentation",
                "Apply to primary lending institutions",
                "Maintain current financial position"
            ],
            "box_type": "success"
        })
    elif score >= 55:
        recommendations.append({
            "title": "Borderline Credit Profile",
            "message": f"Score of {score}/100 requires enhanced documentation.",
            "actions": [
                "Provide detailed income verification",
                "Include employment stability documentation",
                "Consider additional co-signers if applicable"
            ],
            "box_type": "warning"
        })
    else:
        recommendations.append({
            "title": "Profile Requires Strengthening",
            "message": f"Score of {score}/100 indicates significant improvement needed.",
            "actions": [
                "Focus on credit score improvement for 6-12 months",
                "Increase liquid asset position",
                "Consult with financial advisory services"
            ],
            "box_type": "warning"
        })
    
    return recommendations

def generate_shap_explanation(model, X_input, feature_names):
    """Generate SHAP-based explanation for model decision"""
    try:
        # Create SHAP explainer
        explainer = shap.TreeExplainer(model)
        shap_values = explainer.shap_values(X_input)
        
        # For binary classification
        if isinstance(shap_values, list):
            shap_values = shap_values[1]
        
        # Get feature importance
        feature_importance = np.abs(shap_values).mean(0)
        feature_df = pd.DataFrame({
            'Feature': feature_names,
            'Impact': feature_importance[0] if len(feature_importance.shape) > 1 else feature_importance
        }).sort_values('Impact', ascending=False).head(10)
        
        return feature_df
    except Exception as e:
        st.error(f"SHAP explanation error: {str(e)}")
        return None

# ---------------------------------------------------------
# WORD DOCUMENT REPORT GENERATION
# ---------------------------------------------------------
def save_plotly_fig(fig, filename):
    """Save Plotly figure as image"""
    img_bytes = pio.to_image(fig, format='png', width=600, height=400)
    with open(filename, 'wb') as f:
        f.write(img_bytes)
    return filename

def create_word_report(applicant_data, results, features, shap_data):
    """Generate comprehensive Word document report"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Credit Risk Assessment Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f'Generated: {datetime.now().strftime("%d %B %Y at %H:%M")}')
    date_run.italic = True
    
    doc.add_paragraph()
    
    # Applicant Information Section
    doc.add_heading('Applicant Information', level=1)
    
    info_table = doc.add_table(rows=7, cols=2)
    info_table.style = 'LightShading-Accent1'
    
    rows = [
        ("Credit Score:", str(applicant_data['credit_score'])),
        ("Annual Income:", f"£{applicant_data['income_annum']:,}"),
        ("Loan Amount:", f"£{applicant_data['loan_amount']:,}"),
        ("Loan Term:", f"{applicant_data['loan_term']} years"),
        ("Employment:", applicant_data['self_employed']),
        ("Education:", applicant_data['education']),
        ("Total Assets:", f"£{applicant_data['total_assets']:,}")
    ]
    
    for i, (label, value) in enumerate(rows):
        info_table.cell(i, 0).text = label
        info_table.cell(i, 1).text = value
    
    doc.add_paragraph()
    
    # Assessment Results
    doc.add_heading('Assessment Results', level=1)
    
    results_table = doc.add_table(rows=3, cols=2)
    results_table.style = 'LightShading-Accent1'
    
    results_rows = [
        ("Matching Score:", f"{results['matching_score']}/100"),
        ("Approval Probability:", f"{results['approval_probability']:.1f}%"),
        ("Risk Assessment:", results['status'])
    ]
    
    for i, (label, value) in enumerate(results_rows):
        results_table.cell(i, 0).text = label
        results_table.cell(i, 1).text = value
    
    doc.add_paragraph()
    
    # Financial Health Metrics
    doc.add_heading('Financial Health Metrics', level=1)
    
    metrics_table = doc.add_table(rows=4, cols=2)
    metrics_table.style = 'LightShading-Accent1'
    
    metrics_rows = [
        ("Monthly Payment Ratio:", f"{features['monthly_payment_ratio']:.1f}%"),
        ("Asset Coverage:", f"{features['asset_coverage']:.1f}%"),
        ("Credit Category:", features['credit_category']),
        ("Stability Score:", f"{features['stability_score']}/45")
    ]
    
    for i, (label, value) in enumerate(metrics_rows):
        metrics_table.cell(i, 0).text = label
        metrics_table.cell(i, 1).text = value
    
    doc.add_page_break()
    
    # Charts Section
    doc.add_heading('Visual Analysis', level=1)
    
    # Gauge Chart
    doc.add_heading('Matching Score Gauge', level=2)
    gauge_fig = create_score_gauge(results['matching_score'])
    gauge_file = save_plotly_fig(gauge_fig, 'gauge_chart.png')
    doc.add_picture(gauge_file, width=Inches(6))
    doc.add_paragraph(f"Score Interpretation: {results['status']}")
    
    doc.add_paragraph()
    
    # Radar Chart
    doc.add_heading('Financial Health Radar', level=2)
    radar_fig = create_feature_radar(features)
    radar_file = save_plotly_fig(radar_fig, 'radar_chart.png')
    doc.add_picture(radar_file, width=Inches(6))
    
    doc.add_paragraph()
    
    # SHAP Chart
    if shap_data is not None and not shap_data.empty:
        doc.add_heading('Decision Factors Analysis', level=2)
        shap_fig = create_shap_chart(shap_data)
        if shap_fig:
            shap_file = save_plotly_fig(shap_fig, 'shap_chart.png')
            doc.add_picture(shap_file, width=Inches(6))
            
            # Add SHAP explanation
            doc.add_paragraph()
            doc.add_heading('Top Decision Factors', level=3)
            shap_table = doc.add_table(rows=min(6, len(shap_data))+1, cols=2)
            shap_table.style = 'LightShading-Accent1'
            shap_table.cell(0, 0).text = "Feature"
            shap_table.cell(0, 1).text = "Impact Score"
            
            for i, row in shap_data.head(6).iterrows():
                shap_table.cell(i+1, 0).text = row['Feature'].replace('_', ' ').title()
                shap_table.cell(i+1, 1).text = f"{row['Impact']:.4f}"
    
    doc.add_page_break()
    
    # Recommendations Section
    doc.add_heading('Credit Risk Recommendations', level=1)
    
    for i, rec in enumerate(results['recommendations'][:5], 1):
        doc.add_heading(f"{i}. {rec['title']}", level=2)
        doc.add_paragraph(rec['message'])
        
        if rec['actions']:
            doc.add_heading("Recommended Actions:", level=3)
            for action in rec['actions']:
                para = doc.add_paragraph(style='ListBullet')
                para.add_run(action)
        
        doc.add_paragraph()
    
    # Credit Management Advice
    doc.add_heading('Credit Management Guidance', level=1)
    
    advice_items = [
        "Maintain credit utilization below 30% of available limits",
        "Ensure all payments are made on time, every time",
        "Regularly review credit reports for inaccuracies",
        "Avoid multiple credit applications within short periods",
        "Build and maintain emergency savings equivalent to 3-6 months of expenses",
        "Diversify credit types responsibly over time",
        "Monitor debt-to-income ratio monthly",
        "Establish long-term banking relationships"
    ]
    
    for item in advice_items:
        para = doc.add_paragraph(style='ListBullet')
        para.add_run(item)
    
    doc.add_paragraph()
    
    # Disclaimer
    disclaimer = doc.add_paragraph()
    disclaimer_run = disclaimer.add_run('Disclaimer')
    disclaimer_run.bold = True
    
    doc.add_paragraph("""This assessment is generated by machine learning models and provides preliminary risk evaluation only. 
    Final credit decisions are made by individual lending institutions based on comprehensive underwriting criteria. 
    This report does not constitute a loan offer or guarantee of credit approval. 
    All financial decisions should be made in consultation with qualified financial advisors.""")
    
    # Save to BytesIO
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    
    return doc_bytes

# ---------------------------------------------------------
# VISUALIZATION FUNCTIONS
# ---------------------------------------------------------
def create_score_gauge(score):
    """Create a professional gauge chart for matching score"""
    colors = ['#ef5350', '#ff9800', '#4caf50', '#2e7d32']
    
    fig = go.Figure(go.Indicator(
        mode="gauge+number",
        value=score,
        domain={'x': [0, 1], 'y': [0, 1]},
        title={'text': "Credit Profile Score", 'font': {'size': 18}},
        gauge={
            'axis': {'range': [None, 100], 'tickwidth': 1, 'tickcolor': "#37474f"},
            'bar': {'color': "#1a237e"},
            'bgcolor': "rgba(0,0,0,0)",
            'borderwidth': 2,
            'bordercolor': "#546e7a",
            'steps': [
                {'range': [0, 50], 'color': colors[0]},
                {'range': [50, 70], 'color': colors[1]},
                {'range': [70, 85], 'color': colors[2]},
                {'range': [85, 100], 'color': colors[3]}
            ],
            'threshold': {
                'line': {'color': "#000000", 'width': 3},
                'thickness': 0.85,
                'value': 70
            }
        }
    ))
    
    fig.update_layout(
        paper_bgcolor='rgba(0,0,0,0)',
        plot_bgcolor='rgba(0,0,0,0)',
        font=dict(color='#263238', family="Arial"),
        height=300,
        margin=dict(l=30, r=30, t=50, b=30)
    )
    
    return fig

def create_feature_radar(features):
    """Create radar chart for financial health metrics"""
    categories = ['Credit Score', 'Affordability', 'Asset Security', 'Stability']
    values = [
        min(100, features['credit_score'] / 9.99),
        100 - min(100, features['monthly_payment_ratio']),
        min(100, features['asset_coverage'] / 2.5),
        min(100, features['stability_score'] / 45 * 100)
    ]
    
    fig = go.Figure(data=go.Scatterpolar(
        r=values + [values[0]],
        theta=categories + [categories[0]],
        fill='toself',
        fillcolor='rgba(26, 35, 126, 0.2)',
        line_color='rgb(26, 35, 126)',
        line_width=2
    ))
    
    fig.update_layout(
        polar=dict(
            radialaxis=dict(
                visible=True,
                range=[0, 100],
                gridcolor='rgba(96, 125, 139, 0.3)',
                linecolor='#546e7a'
            ),
            angularaxis=dict(
                gridcolor='rgba(96, 125, 139, 0.3)',
                linecolor='#546e7a'
            ),
            bgcolor='rgba(0,0,0,0)'
        ),
        paper_bgcolor='rgba(0,0,0,0)',
        font=dict(color='#263238'),
        showlegend=False,
        height=300,
        margin=dict(l=50, r=50, t=30, b=30)
    )
    
    return fig

def create_shap_chart(shap_data):
    """Create horizontal bar chart for SHAP feature importance"""
    if shap_data is None or shap_data.empty:
        return None
    
    plot_data = shap_data.head(8).sort_values('Impact', ascending=True)
    
    fig = go.Figure(go.Bar(
        x=plot_data['Impact'],
        y=plot_data['Feature'].str.replace('_', ' ').str.title(),
        orientation='h',
        marker_color='#1a237e',
        text=plot_data['Impact'].round(4),
        textposition='outside'
    ))
    
    fig.update_layout(
        title='Feature Impact Analysis',
        xaxis_title='Impact Magnitude',
        yaxis_title='Risk Factors',
        paper_bgcolor='rgba(0,0,0,0)',
        plot_bgcolor='rgba(0,0,0,0)',
        font=dict(color='#263238'),
        height=300,
        margin=dict(l=150, r=20, t=40, b=20)
    )
    
    return fig
def save_plotly_fig(fig, filename):
    """Save Plotly figure as image with Kaleido configuration"""
    try:
        # Configure Kaleido for headless environments (critical for Streamlit Cloud)
        import plotly.io as pio
        
        # Set Chromium arguments for headless environments[citation:1]
        if hasattr(pio.kaleido.scope, 'chromium_args'):
            pio.kaleido.scope.chromium_args = (
                "--headless",
                "--no-sandbox", 
                "--single-process",
                "--disable-gpu",
                "--disable-dev-shm-usage"
            )
        
        # Save the figure
        img_bytes = pio.to_image(fig, format='png', width=600, height=400)
        with open(filename, 'wb') as f:
            f.write(img_bytes)
        return filename
        
    except Exception as e:
        # Fallback: Save as HTML if PNG export fails
        st.warning(f"Image export using Kaleido failed: {str(e)}. Using HTML fallback.")
        html_filename = filename.replace('.png', '.html')
        pio.write_html(fig, file=html_filename)
        return html_filename
# ---------------------------------------------------------
# MAIN APPLICATION
# ---------------------------------------------------------
def main():
    # Header
    st.markdown('<h1 class="main-header">Credit Risk Assessment System</h1>', unsafe_allow_html=True)
    st.markdown("""
    <div class="info-container">
    <strong>Professional Risk Evaluation:</strong> This system provides data-driven credit risk assessment using machine learning models. 
    All evaluations are based on UK lending standards and risk management principles.
    </div>
    """, unsafe_allow_html=True)
    
    # Sidebar for inputs
    with st.sidebar:
        st.markdown('<div class="sidebar-section">', unsafe_allow_html=True)
        st.markdown('<h3>Application Information</h3>', unsafe_allow_html=True)
        
        # Personal Information
        st.subheader("Personal Details")
        credit_score = st.slider(
            "Credit Score (UK Scale)",
            min_value=0,
            max_value=999,
            value=750,
            help="UK credit scores range from 0-999"
        )
        
        income_annum = st.number_input(
            "Annual Income (£)",
            min_value=15000,
            value=35000,
            step=5000,
            help="Gross annual income before tax"
        )
        
        loan_amount = st.number_input(
            "Loan Amount (£)",
            min_value=5000,
            value=25000,
            step=5000,
            help="Total borrowing requirement"
        )
        
        loan_term = st.slider(
            "Loan Term (Years)",
            min_value=1,
            max_value=30,
            value=5,
            help="Repayment period in years"
        )
        
        no_of_dependents = st.selectbox(
            "Number of Dependents",
            options=[0, 1, 2, 3, 4, "5 or more"],
            index=1
        )
        
        # Employment & Education
        col1, col2 = st.columns(2)
        with col1:
            self_employed = st.radio(
                "Employment Type",
                ["Employed", "Self-Employed"]
            )
        with col2:
            education = st.radio(
                "Education Level",
                ["Graduate", "Not Graduate"]
            )
        
        st.markdown('</div>', unsafe_allow_html=True)
        
        # Asset Information
        st.markdown('<div class="sidebar-section">', unsafe_allow_html=True)
        st.markdown('<h3>Asset Information</h3>', unsafe_allow_html=True)
        
        total_assets = st.number_input(
            "Total Assets Value (£)",
            min_value=0,
            value=50000,
            step=10000,
            help="Combined value of all assets"
        )
        
        if st.checkbox("View asset allocation model"):
            st.info("""
            **Standard Asset Allocation:**
            - Residential Assets: 50%
            - Commercial Assets: 25%
            - Luxury Assets: 15%
            - Bank Assets: 10%
            """)
        st.markdown('</div>', unsafe_allow_html=True)
        
        # Rate limiting
        current_time = time.time()
        time_since_last = current_time - st.session_state.last_submission
        
        if time_since_last < 3 and st.session_state.results is not None:
            remaining = 3 - time_since_last
            st.warning(f"Assessment cooling period: {remaining:.1f} seconds")
            calculate_disabled = True
        else:
            calculate_disabled = False
        
        # Calculate Button
        calculate_btn = st.button(
            "Execute Risk Assessment",
            type="primary",
            disabled=calculate_disabled,
            use_container_width=True
        )
    
    # Main Content Area
    if calculate_btn:
        # Update last submission time
        st.session_state.last_submission = time.time()
        
        # Validate inputs
        validation_errors = []
        if loan_term > 30:
            validation_errors.append("Maximum loan term is 30 years")
        if income_annum <= 0:
            validation_errors.append("Annual income must be positive")
        if not (0 <= credit_score <= 999):
            validation_errors.append("Credit score must be between 0-999")
        
        if validation_errors:
            for error in validation_errors:
                st.error(f"Validation Error: {error}")
            return
        
        # Asset allocation
        ASSET_ALLOCATION = {
            "residential_assets_value": 0.50,
            "commercial_assets_value": 0.25,
            "luxury_assets_value": 0.15,
            "bank_asset_value": 0.10
        }
        
        # Convert to model scale
        residential_assets = int((total_assets * ASSET_ALLOCATION['residential_assets_value']) * CONVERSION_RATE)
        commercial_assets = int((total_assets * ASSET_ALLOCATION['commercial_assets_value']) * CONVERSION_RATE)
        luxury_assets = int((total_assets * ASSET_ALLOCATION['luxury_assets_value']) * CONVERSION_RATE)
        bank_assets = int((total_assets * ASSET_ALLOCATION['bank_asset_value']) * CONVERSION_RATE)
        
        # Prepare input data
        model_input_data = {
            "credit_score": credit_score,
            "income_annum": int(income_annum * CONVERSION_RATE),
            "loan_amount": int(loan_amount * CONVERSION_RATE),
            "loan_term": loan_term,
            "no_of_dependents": no_of_dependents if isinstance(no_of_dependents, int) else 5,
            "self_employed": "Yes" if self_employed == "Self-Employed" else "No",
            "education": education,
            "residential_assets_value": residential_assets,
            "commercial_assets_value": commercial_assets,
            "luxury_assets_value": luxury_assets,
            "bank_asset_value": bank_assets,
            "total_assets": int(total_assets * CONVERSION_RATE)
        }
        
        # Display data
        uk_display_data = {
            "credit_score": credit_score,
            "income_annum": income_annum,
            "loan_amount": loan_amount,
            "loan_term": loan_term,
            "no_of_dependents": no_of_dependents if isinstance(no_of_dependents, int) else 5,
            "self_employed": self_employed,
            "education": education,
            "total_assets": total_assets
        }
        
        # Create DataFrame for model
        df_input = pd.DataFrame([model_input_data])
        
        # Feature engineering
        df_features = create_credit_features(df_input)
        
        # Calculate matching score
        matching_score = calculate_matching_score(df_features.iloc[0])
        
        # Load model and predict - NO FALLBACK
        try:
            model = joblib.load("best_xgb_model.pkl")
            feature_columns = joblib.load("feature_columns.pkl")
            
            # Prepare features for model
            X_input = df_input.drop(columns=['total_assets'])
            X_input = pd.get_dummies(X_input)
            X_input = X_input.reindex(columns=feature_columns, fill_value=0)
            
            # Get prediction
            approval_probability = model.predict_proba(X_input)[0, 1] * 100
            prediction = model.predict(X_input)[0]
            
            # Generate SHAP explanation
            shap_data = generate_shap_explanation(model, X_input, feature_columns)
            
        except FileNotFoundError as e:
            st.error(f"Model file not found: {e}")
            st.stop()
        except Exception as e:
            st.error(f"Model execution error: {str(e)}")
            st.stop()
        
        # Generate detailed recommendations
        recommendations = generate_detailed_recommendations(
            matching_score, 
            df_features.iloc[0].to_dict(),
            uk_display_data,
            prediction
        )
        
        # Determine status
        if matching_score >= 80:
            status = "Low Risk Profile"
            status_box = "success"
        elif matching_score >= 65:
            status = "Moderate Risk Profile"
            status_box = "success"
        elif matching_score >= 50:
            status = "Elevated Risk Profile"
            status_box = "warning"
        else:
            status = "High Risk Profile"
            status_box = "warning"
        
        # Store results
        st.session_state.results = {
            'matching_score': matching_score,
            'approval_probability': approval_probability,
            'prediction': prediction,
            'status': status,
            'status_box': status_box,
            'recommendations': recommendations,
            'applicant_data': uk_display_data,
            'features': df_features.iloc[0].to_dict(),
            'shap_data': shap_data
        }
    
    # Display results if available
    if st.session_state.results:
        results = st.session_state.results
        
        # Results Header
        st.markdown('<h2 class="sub-header">Risk Assessment Report</h2>', unsafe_allow_html=True)
        
        # Key Metrics
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric(
                label="Profile Score",
                value=f"{results['matching_score']}/100",
                delta=results['status']
            )
        
        with col2:
            st.metric(
                label="Approval Probability",
                value=f"{results['approval_probability']:.1f}%"
            )
        
        with col3:
            decision = "Approve" if results['prediction'] == 1 else "Review"
            st.metric(
                label="Model Decision",
                value=decision
            )
        
        with col4:
            if results['matching_score'] >= 70:
                time_value = "Standard"
            elif results['matching_score'] >= 50:
                time_value = "Extended"
            else:
                time_value = "Manual"
            st.metric(
                label="Processing",
                value=time_value
            )
        
        # Charts
        col1, col2 = st.columns(2)
        
        with col1:
            st.plotly_chart(create_score_gauge(results['matching_score']), use_container_width=True)
        
        with col2:
            st.plotly_chart(create_feature_radar(results['features']), use_container_width=True)
        
        # SHAP Explanation Chart
        if results.get('shap_data') is not None:
            shap_chart = create_shap_chart(results['shap_data'])
            if shap_chart:
                st.plotly_chart(shap_chart, use_container_width=True)
        
        # Financial Health Indicators
        st.markdown('<h3 class="sub-header">Financial Health Metrics</h3>', unsafe_allow_html=True)
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            payment_ratio = results['features']['monthly_payment_ratio']
            st.write("**Payment Affordability**")
            progress_value = max(0.0, min(1.0, (35 - min(payment_ratio, 35)) / 35))
            st.progress(
                progress_value,
                text=f"{payment_ratio:.1f}% of income"
            )
            st.caption("Target: ≤35% of monthly income")
        
        with col2:
            asset_coverage = results['features']['asset_coverage']
            st.write("**Asset Coverage**")
            progress_value = min(1.0, asset_coverage / 250)
            st.progress(
                progress_value,
                text=f"{asset_coverage:.1f}% coverage"
            )
            st.caption("Ideal: ≥125% of loan amount")
        
        with col3:
            stability = results['features']['stability_score']
            st.write("**Stability Score**")
            progress_value = min(1.0, stability / 60)
            st.progress(
                progress_value,
                text=f"{stability}/60 points"
            )
            st.caption("Employment, education, dependents")
        
        # Risk Recommendations
        st.markdown('<h3 class="sub-header">Risk Mitigation Recommendations</h3>', unsafe_allow_html=True)
        
        for i, rec in enumerate(results['recommendations'][:5]):
            if rec['box_type'] == "success":
                box_class = "success-container"
            elif rec['box_type'] == "warning":
                box_class = "warning-container"
            else:
                box_class = "info-container"
            
            with st.expander(f"{rec['title']}", expanded=i<2):
                st.markdown(f"<div class='{box_class}'>", unsafe_allow_html=True)
                st.write(f"**{rec['message']}**")
                st.write("")
                st.write("**Action Items:**")
                for action in rec['actions']:
                    st.write(f"• {action}")
                st.markdown("</div>", unsafe_allow_html=True)
        
        # Word Document Generation
        st.markdown('<h3 class="sub-header">Download Assessment Report</h3>', unsafe_allow_html=True)
        
        col1, col2 = st.columns([3, 1])
        
        with col1:
            st.info("Generate a comprehensive Word document containing all assessment details, visualizations, and credit management recommendations.")
        
        with col2:
            try:
                doc_bytes = create_word_report(
                    results['applicant_data'],
                    results,
                    results['features'],
                    results.get('shap_data')
                )
                
                b64 = base64.b64encode(doc_bytes.getvalue()).decode()
                current_date = datetime.now().strftime("%Y%m%d")
                href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="Credit_Assessment_{current_date}.docx" style="background-color: #1a237e; color: white; padding: 12px 24px; text-decoration: none; border-radius: 6px; display: inline-block; font-weight: 600;">Download Word Report</a>'
                st.markdown(href, unsafe_allow_html=True)
            except Exception as e:
                st.error(f"Report generation failed: {str(e)}")
    
    else:
        # Welcome screen
        col1, col2 = st.columns([2, 1])
        
        with col1:
            st.markdown("""
            ### Assessment Methodology
            
            **1. Data Collection**
            Comprehensive collection of financial and personal information using standardized input forms.
            
            **2. Machine Learning Analysis**
            Advanced XGBoost model processes 12+ risk factors using historical lending data.
            
            **3. Risk Scoring**
            Generation of 0-100 risk score based on creditworthiness assessment.
            
            **4. SHAP Explanation**
            Transparent explanation of model decisions using Shapley values.
            
            **5. Actionable Reporting**
            Detailed recommendations and professional documentation.
            
            ### Risk Assessment Framework
            
            • **Credit History**: Analysis of credit score and payment patterns  
            • **Payment Capacity**: Evaluation of debt-to-income ratios  
            • **Asset Security**: Assessment of collateral and liquid assets  
            • **Employment Stability**: Review of income consistency and source  
            • **Demographic Factors**: Consideration of education and dependents  
            
            ### System Capabilities
            
            • **Predictive Accuracy**: Machine learning model trained on historical data  
            • **Transparent Decisions**: SHAP-based feature importance explanations  
            • **Professional Reporting**: Comprehensive Word document generation  
            • **Real-time Processing**: Instant assessment with rate limiting  
            • **UK Compliance**: Adherence to UK lending standards and regulations  
            """)
        
        with col2:
            st.markdown("""
            <div class="info-container">
            <h4>UK Credit Standards</h4>
            <ul style="padding-left: 1.2rem; margin-bottom: 0;">
            <li><strong>Credit Score ≥700</strong>: Preferred lending criteria</li>
            <li><strong>Monthly payments ≤35%</strong>: Recommended affordability threshold</li>
            <li><strong>Asset coverage ≥125%</strong>: Optimal security position</li>
            <li><strong>Employment history ≥2 years</strong>: Stability benchmark</li>
            <li><strong>Debt-to-income ≤40%</strong>: Maximum recommended ratio</li>
            <li><strong>Credit inquiries ≤3</strong>: Annual application limit</li>
            </ul>
            </div>
            """, unsafe_allow_html=True)
    
    # Footer with Disclaimer
    st.markdown("---")
    st.markdown("""
    <div class="disclaimer-box">
    <p><strong>Important Notice:</strong> This credit risk assessment system provides preliminary evaluation based on machine learning models. All assessments are indicative and do not constitute financial advice or credit guarantees. Final lending decisions remain at the discretion of individual financial institutions based on comprehensive underwriting processes. Users should consult with qualified financial advisors before making borrowing decisions. Model accuracy may vary based on input data quality and completeness.</p>
    <p style="margin-top: 0.5rem; text-align: center; font-weight: 600;">© 2026 Smart Solution to tough data</p>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()
